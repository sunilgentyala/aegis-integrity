"""
Multi-format academic document parser.
Supports: PDF, DOCX, LaTeX (.tex), plain text, BibTeX.
"""

from __future__ import annotations
import re
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class ParsedSection:
    title: str
    level: int          # 1=section, 2=subsection, 3=subsubsection
    text: str
    start_char: int
    end_char: int


@dataclass
class ParsedReference:
    raw: str
    authors: list[str]
    year: Optional[str]
    title: Optional[str]
    doi: Optional[str]
    url: Optional[str]
    journal: Optional[str]
    cite_key: Optional[str]     # BibTeX key or in-text marker
    line_number: Optional[int]


@dataclass
class ParsedDocument:
    path: str
    format: str                          # pdf | docx | latex | txt | bib
    title: Optional[str]
    authors: list[str]
    abstract: Optional[str]
    full_text: str
    sections: list[ParsedSection]
    references: list[ParsedReference]
    metadata: dict = field(default_factory=dict)

    @property
    def body_text(self) -> str:
        """Full text excluding references section."""
        if not self.sections:
            return self.full_text
        body_sections = [s for s in self.sections
                         if "reference" not in s.title.lower()
                         and "bibliograph" not in s.title.lower()]
        return "\n\n".join(s.text for s in body_sections) or self.full_text

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


class DocumentParser:
    """Parse academic documents into structured form."""

    def parse(self, path: str | Path) -> ParsedDocument:
        path = Path(path)
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext == ".docx":
            return self._parse_docx(path)
        elif ext in (".tex",):
            return self._parse_latex(path)
        elif ext == ".bib":
            return self._parse_bibtex(path)
        else:
            return self._parse_txt(path)

    # ------------------------------------------------------------------
    def _parse_pdf(self, path: Path) -> ParsedDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF required: pip install PyMuPDF")

        doc = fitz.open(str(path))
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text("text"))
        full_text = "\n".join(pages_text)
        doc.close()

        sections = self._extract_sections_heuristic(full_text)
        abstract = self._extract_abstract(full_text)
        title = self._extract_title_heuristic(full_text)
        refs = self._extract_references_heuristic(full_text)

        return ParsedDocument(
            path=str(path), format="pdf",
            title=title, authors=[],
            abstract=abstract, full_text=full_text,
            sections=sections, references=refs,
            metadata={"pages": len(pages_text)},
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paras)
        sections = self._extract_sections_heuristic(full_text)
        abstract = self._extract_abstract(full_text)
        title = paras[0] if paras else None
        refs = self._extract_references_heuristic(full_text)

        return ParsedDocument(
            path=str(path), format="docx",
            title=title, authors=[],
            abstract=abstract, full_text=full_text,
            sections=sections, references=refs,
        )

    def _parse_latex(self, path: Path) -> ParsedDocument:
        raw = path.read_text(encoding="utf-8", errors="replace")

        # Extract title
        title_m = re.search(r"\\title\{([^}]+)\}", raw, re.DOTALL)
        title = title_m.group(1).strip() if title_m else None

        # Extract authors
        authors = re.findall(r"\\author(?:\[.*?\])?\{([^}]+)\}", raw)

        # Extract abstract
        abs_m = re.search(r"\\begin\{abstract\}(.+?)\\end\{abstract\}", raw,
                          re.DOTALL)
        abstract = self._clean_latex(abs_m.group(1)) if abs_m else None

        # Extract sections
        sections = self._extract_latex_sections(raw)

        # Full text (cleaned)
        full_text = self._clean_latex(raw)

        # References from \bibliography or inline bibitem
        refs = self._extract_latex_references(raw)

        return ParsedDocument(
            path=str(path), format="latex",
            title=self._clean_latex(title) if title else None,
            authors=authors,
            abstract=abstract, full_text=full_text,
            sections=sections, references=refs,
        )

    def _parse_bibtex(self, path: Path) -> ParsedDocument:
        raw = path.read_text(encoding="utf-8", errors="replace")
        refs = self._parse_bibtex_entries(raw)
        return ParsedDocument(
            path=str(path), format="bib",
            title=None, authors=[], abstract=None,
            full_text=raw, sections=[], references=refs,
        )

    def _parse_txt(self, path: Path) -> ParsedDocument:
        full_text = path.read_text(encoding="utf-8", errors="replace")
        sections = self._extract_sections_heuristic(full_text)
        abstract = self._extract_abstract(full_text)
        refs = self._extract_references_heuristic(full_text)
        return ParsedDocument(
            path=str(path), format="txt",
            title=None, authors=[], abstract=abstract,
            full_text=full_text, sections=sections, references=refs,
        )

    # ------------------------------------------------------------------
    # LaTeX helpers
    # ------------------------------------------------------------------

    def _clean_latex(self, text: str) -> str:
        if not text:
            return ""
        # Remove comments
        text = re.sub(r"%.*", "", text)
        # Remove common commands but keep their argument text
        text = re.sub(r"\\(?:textbf|textit|emph|underline|texttt|textrm)\{([^}]*)\}",
                      r"\1", text)
        # Remove \cite, \ref, \label
        text = re.sub(r"\\(?:cite|ref|label|eqref)\{[^}]*\}", "", text)
        # Remove \begin{...}...\end{...} for math/table/figure
        text = re.sub(r"\\begin\{(?:equation|align|table|figure|tabular|array)[*]?\}.*?"
                      r"\\end\{(?:equation|align|table|figure|tabular|array)[*]?\}",
                      "", text, flags=re.DOTALL)
        # Remove remaining commands
        text = re.sub(r"\\[a-zA-Z]+\*?\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\[a-zA-Z]+\*?", "", text)
        text = re.sub(r"[{}]", "", text)
        return re.sub(r"\s+", " ", text).strip()

    def _extract_latex_sections(self, raw: str) -> list[ParsedSection]:
        pattern = re.compile(
            r"\\(section|subsection|subsubsection)\*?\{([^}]+)\}", re.DOTALL)
        level_map = {"section": 1, "subsection": 2, "subsubsection": 3}
        matches = list(pattern.finditer(raw))
        sections = []
        for i, m in enumerate(matches):
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(raw)
            body = self._clean_latex(raw[start:end])
            sections.append(ParsedSection(
                title=self._clean_latex(m.group(2)),
                level=level_map[m.group(1)],
                text=body,
                start_char=start,
                end_char=end,
            ))
        return sections

    def _extract_latex_references(self, raw: str) -> list[ParsedReference]:
        refs = []
        # \bibitem{key} ... text
        bibitem_pattern = re.compile(r"\\bibitem(?:\[.*?\])?\{([^}]+)\}(.*?)(?=\\bibitem|\\end\{thebibliography\})",
                                     re.DOTALL)
        for m in bibitem_pattern.finditer(raw):
            key = m.group(1).strip()
            text = self._clean_latex(m.group(2)).strip()
            doi = self._extract_doi(text)
            refs.append(ParsedReference(
                raw=text, authors=[], year=self._extract_year(text),
                title=None, doi=doi, url=None, journal=None,
                cite_key=key, line_number=None,
            ))
        return refs

    def _parse_bibtex_entries(self, raw: str) -> list[ParsedReference]:
        refs = []
        entry_pattern = re.compile(
            r"@\w+\{(\w+),([^@]+)", re.DOTALL)
        for m in entry_pattern.finditer(raw):
            key = m.group(1)
            body = m.group(2)
            doi   = self._field(body, "doi")
            title = self._field(body, "title")
            year  = self._field(body, "year")
            auth  = self._field(body, "author")
            jour  = self._field(body, "journal") or self._field(body, "booktitle")
            authors = [a.strip() for a in auth.split(" and ")] if auth else []
            refs.append(ParsedReference(
                raw=body, authors=authors, year=year, title=title,
                doi=doi, url=self._field(body, "url"),
                journal=jour, cite_key=key, line_number=None,
            ))
        return refs

    def _field(self, bib_body: str, name: str) -> Optional[str]:
        m = re.search(rf"{name}\s*=\s*\{{([^}}]*)\}}", bib_body, re.IGNORECASE)
        if m:
            return m.group(1).strip()
        m = re.search(rf'{name}\s*=\s*"([^"]*)"', bib_body, re.IGNORECASE)
        return m.group(1).strip() if m else None

    # ------------------------------------------------------------------
    # Generic heuristic helpers
    # ------------------------------------------------------------------

    def _extract_sections_heuristic(self, text: str) -> list[ParsedSection]:
        """Detect section breaks from ALL-CAPS lines or numbered headings."""
        sections = []
        lines = text.split("\n")
        current_title = "Introduction"
        current_lines: list[str] = []
        start_char = 0
        char_pos = 0

        heading_pattern = re.compile(
            r"^(?:\d+\.?\s+)?([A-Z][A-Z\s&:/\-]{3,60})$")

        for line in lines:
            stripped = line.strip()
            if heading_pattern.match(stripped) and len(stripped) < 80:
                if current_lines:
                    body = "\n".join(current_lines).strip()
                    end_char = char_pos
                    sections.append(ParsedSection(
                        title=current_title, level=1,
                        text=body, start_char=start_char, end_char=end_char,
                    ))
                current_title = stripped
                current_lines = []
                start_char = char_pos
            else:
                current_lines.append(line)
            char_pos += len(line) + 1

        if current_lines:
            sections.append(ParsedSection(
                title=current_title, level=1,
                text="\n".join(current_lines).strip(),
                start_char=start_char, end_char=char_pos,
            ))
        return sections or [ParsedSection("Body", 1, text, 0, len(text))]

    def _extract_abstract(self, text: str) -> Optional[str]:
        m = re.search(
            r"(?:Abstract|ABSTRACT)[:\s]*\n(.*?)(?:\n\n|\n[A-Z]{3,})",
            text, re.DOTALL)
        if m:
            return m.group(1).strip()
        return None

    def _extract_title_heuristic(self, text: str) -> Optional[str]:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:10]:
            if 10 < len(line) < 200 and not line.startswith("http"):
                return line
        return None

    def _extract_references_heuristic(self, text: str) -> list[ParsedReference]:
        refs = []
        # Find reference section
        ref_match = re.search(
            r"(?:References|Bibliography|REFERENCES)\s*\n(.*)",
            text, re.DOTALL | re.IGNORECASE)
        if not ref_match:
            return []
        ref_block = ref_match.group(1)
        # Split on numbered entries [1] or (1) or line starts
        entries = re.split(r"\n\s*(?:\[\d+\]|\(\d+\)|(?:\d+\.))\s+", ref_block)
        for i, entry in enumerate(entries):
            entry = entry.strip()
            if len(entry) < 20:
                continue
            doi = self._extract_doi(entry)
            refs.append(ParsedReference(
                raw=entry, authors=[],
                year=self._extract_year(entry),
                title=None, doi=doi, url=None, journal=None,
                cite_key=f"ref_{i}", line_number=None,
            ))
        return refs

    def _extract_doi(self, text: str) -> Optional[str]:
        m = re.search(r"10\.\d{4,}/\S+", text)
        return m.group(0).rstrip(".,;)") if m else None

    def _extract_year(self, text: str) -> Optional[str]:
        m = re.search(r"\b(19|20)\d{2}\b", text)
        return m.group(0) if m else None
